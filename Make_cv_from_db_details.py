import PyPDF2
import spacy
import docx
from docx2pdf import convert
import pyrebase
import os
import webbrowser

def make_cv(lists):
        doc = docx.Document(r"cv_template.docx")
        # Find the text you want to replace
        old_text = "desc"
        for paragraph in doc.paragraphs:
            if old_text in paragraph.text:
                # Replace the text while maintaining the same style and formatting
                for run in paragraph.runs:
                    if old_text in run.text:
                        new_text = run.text.replace(old_text, lists['job_desc_list'][0][0])
                        run.text = new_text
                        break  # break out of the loop once the text is replaced in one run object
        old_text = "Company"
        for paragraph in doc.paragraphs:
            if old_text in paragraph.text:
                # Replace the text while maintaining the same style and formatting
                for run in paragraph.runs:
                    if old_text in run.text:
                        new_text = run.text.replace(old_text,
                                                    str(lists['companylist'][0][0]) + "|"   +lists['job_fromdatelist'][0][0] + "|" + lists['job_todatelist'][0][0])
                        run.text = new_text
                        break  # break out of the loop once the text is replaced in one run object
        old_text = "Another"
        for paragraph in doc.paragraphs:
            if old_text in paragraph.text:
                # Replace the text while maintaining the same style and formatting
                for run in paragraph.runs:
                    if old_text in run.text:
                        new_text = run.text.replace(old_text,
                                                    str(lists['companylist'][0][1]) + "|" +
                                                    lists['job_fromdatelist'][0][1] + "|" +
                                                    lists['job_todatelist'][0][1])
                        run.text = new_text
                        break  # break out of the loop once the text is replaced in one run object

        old_text = "Marketing"
        for paragraph in doc.paragraphs:
            if old_text in paragraph.text:
                # Replace the text while maintaining the same style and formatting
                for run in paragraph.runs:
                    if old_text in run.text:
                        new_text = run.text.replace(old_text, lists['job_desc_list'][0][1])
                        run.text = new_text
                        break  # break out of the loop once the text is replaced in one run object
        # Save the document
        doc.save("modified3.docx")

        convert(r'modified3.docx', r"modified3.pdf")

        os.remove('modified3.docx')
def make_cv_opti(lists):

    text_dict = {
        'Company':  str(lists['companylist'][0][0]) + "|" + lists['job_fromdatelist'][0][0] + "|" +
                   lists['job_todatelist'][0][0],
        'desc': lists['job_desc_list'][0][0],
        'Another': str(lists['companylist'][0][1]) + "|" + lists['job_fromdatelist'][0][1] + "|" +
                   lists['job_todatelist'][0][1],
        'Marketing': lists['job_desc_list'][0][1],

        'Yet': str(lists['companylist'][0][2]) + "|" + lists['job_fromdatelist'][0][2] + "|" +
                   lists['job_todatelist'][0][2],
        'bro':  lists['job_desc_list'][0][2],
        'next': str(lists['companylist'][0][3]) + "|" + lists['job_fromdatelist'][0][3] + "|" +
                   lists['job_todatelist'][0][3],
       'high' : lists['job_desc_list'][0][3],
        'numba': str(lists['companylist'][0][4]) + "|" + lists['job_fromdatelist'][0][4] + "|" +
        lists['job_todatelist'][0][4],
        'extra': lists['job_desc_list'][0][4],
    }

    doc = docx.Document(r"cv_template - Copy.docx")
    for old_text, new_text in text_dict.items():
        for paragraph in doc.paragraphs:
            if old_text in paragraph.text:
                for run in paragraph.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
                        break

    doc.save("modified4.docx")
    convert(r'modified4.docx', r"modified4.pdf")
    os.remove('modified4.docx')
def dyanamicss(lists):
    # Load the template file
    template = docx.Document('cv_new - Copy.docx')

    # Create a dictionary of placeholders to values
    text_dict = {}
    num_jobs = len(lists['companylist'])
    print(lists['companylist'])
    print(lists['job_desc_list'][0])
    for i in range(num_jobs):
        key1 = f'Company{i + 1}'
        value1 = lists['companylist'][0][i]
        text_dict[key1] = value1

        key2 = f'desc{i + 1}'
        value2 = lists['job_desc_list'][0][i]
        text_dict[key2] = value2
    print(text_dict)

    # Replace the placeholders with values
    for p in template.paragraphs:
        for key, value in text_dict.items():
            if key in p.text:
                inline = p.runs
                # Loop added to work with runs (strings with same style)
                for i in range(len(inline)):
                    if key in inline[i].text:
                        text = inline[i].text.replace(key, value)
                        inline[i].text = text

    # Save the modified template as a new file
    template.save('new_cv_generated.docx')

config = {
    'apiKey': "AIzaSyArHFLsk0pevFNMMGjVavppb5wltduzggw",
    'authDomain': "job-portal-c65be.firebaseapp.com",
    'projectId': "job-portal-c65be",
    'storageBucket': "job-portal-c65be.appspot.com",
    'messagingSenderId': "839506484975",
    'appId': "1:839506484975:web:a703ed5cef40c22f93451a",
    'measurementId': "G-584JGPMYSY",
    'databaseURL': "https://job-portal-c65be-default-rtdb.europe-west1.firebasedatabase.app/",
}
# Initialize the Firebase Storage client
firebase = pyrebase.initialize_app(config)

db = firebase.database()


applicants_data = db.child("Jobs").child("Job Detail").shallow().get().val()
storage = firebase.storage()

applicants=[]
job_id_lis=[]
# Iterate over all the job IDs
for job_id in applicants_data:
    # Get the applicants data for the current job
    data= db.child("Jobs").child("Job Detail").child(job_id).child("Applicants").get().val()
    if data!=None:
        applicants.append(data)
        job_id_lis.append(job_id)

applicants_list = []
# Extract all keys
keys = set()
for item in applicants:
    for d in item.values():
        keys.update(d.keys())

# Initialize empty lists for each key
lists = {key: [] for key in keys}

# Populate the lists
for item in applicants:
    for d in item.values():
        for key, value in d.items():
            lists[key].append(value)

# Print the lists
for key, value in lists.items():
    print(f"{key}: {value}")




companylist=lists['companylist'][0]
companylist
jdlist=lists['job_desc_list'][0]
jdlist
fromlist= lists['job_fromdatelist'][0]
fromlist
tolist=lists['job_todatelist'][0]
tolist
loclist= lists['loc'][0]
loclist
titlelist= lists['Job Title'][0]
titlelist

template = docx.Document('First Name.docx')
paragraphs = template.paragraphs
import re

#
# for p in paragraphs:
#     text = p.text
#     print(text)
#     for i in range(len(companylist)):
#         text = text.replace(f'Company{i+1}', companylist[i])
#         text = text.replace(f'jd{i+1}', jdlist[i])
#         text = text.replace(f'loc{i+1}', loclist[i])
#         text = text.replace(f'from{i+1}', fromlist[i])
#         text = text.replace(f'to{i+1}', tolist[i])
#         text = text.replace(f'Title{i+1}', titlelist[i])
#
#
#     # placeholders = ['Company', 'jd', 'loc', 'from', 'to', 'Title']
#     # for placeholder in placeholders:
#     #     pattern = re.compile(f"{placeholder}\d+")
#     #     text = re.sub(pattern, '', text)
#     # text= text.replace('/', '')
#     # text= text.replace(',', '')
#     # text= text.replace('-', '')
#     #
#     p.text = text
#     template.save('modified_cv2.docx')
#     convert(r'modified_cv1.docx', r"modified_cv2.pdf")
#     webbrowser.open_new_tab(r"modified_cv2.pdf")
#     os.remove('modified_cv2.docx')
#     # print(text)
#
# # Define a list of placeholder names
#
#
# # Save the modified CV

import random
import string

def generate_random_string(length):
    letters = string.ascii_lowercase
    return ''.join(random.choice(letters) for i in range(length))

def uploaditcomplete():
    applicants_data = db.child("Jobs").child("Job Detail").shallow().get().val()
    storage = firebase.storage()

    job_id_lis=[]
    # Iterate over all the job IDs
    for job_id in applicants_data:
        # Get the applicants data for the current job
        data= db.child("Jobs").child("Job Detail").child(job_id).child("Applicants").get().val()
        if data!=None:
            job_id_lis.append(job_id)
            for applicant_id, applicant_data in data.items():
                template = docx.Document('First Name.docx')
                paragraphs = template.paragraphs

                companylist=applicant_data.get('companylist', [])
                jdlist=applicant_data.get('job_desc_list', [])
                fromlist=applicant_data.get('job_fromdatelist', [])
                tolist=applicant_data.get('job_todatelist', [])
                loclist=applicant_data.get('loc', [])
                titlelist=applicant_data.get('Job Title', [])

                for p in paragraphs:
                    text = p.text
                    for i in range(len(companylist)):
                        text = text.replace(f'Company{i+1}', companylist[i])
                        text = text.replace(f'jd{i+1}', jdlist[i])
                        text = text.replace(f'loc{i+1}', loclist[i])
                        text = text.replace(f'from{i+1}', fromlist[i])
                        text = text.replace(f'to{i+1}', tolist[i])
                        text = text.replace(f'Title{i+1}', titlelist[i])
                    placeholders = ['Company', 'jd', 'loc', 'from', 'to', 'Title']
                    for placeholder in placeholders:
                        pattern = re.compile(f"{placeholder}\d+")
                        text = re.sub(pattern, '', text)
                    text= text.replace('/', '')
                    text= text.replace(',', '')
                    p.text = text

                # Save the modified CV
                filename = f"{job_id}_{applicant_id}_{generate_random_string(6)}.docx"
                template.save(filename)
                convert(filename, f"{filename[:-4]}.pdf")

                myfile = f"{filename[:-4]}.pdf"

                storage.child('files/' + myfile).put(myfile)

                url = storage.child('files/' + myfile).get_url(token=None)

                db.child('Jobs').child('Job Detail').child(job_id).child('Applicants').child(applicant_id).update({'template_cv' : url})
                print(f"CV for applicant {applicant_id} in job {job_id} uploaded successfully: {url}")

                # Remove the modified CV
                os.remove(filename)
                os.remove(myfile)

uploaditcomplete()
