import spacy
import docx
from docx2pdf import convert
import os

import nltk
# nltk.download('averaged_perceptron_tagger')
# nltk.downloader.download('maxent_ne_chunker')
# nltk.downloader.download('words')
# nltk.downloader.download('treebank')
# nltk.downloader.download('maxent_treebank_pos_tagger')
# nltk.downloader.download('punkt')
# nltk.download('averaged_perceptron_tagger')
ocr_text = {
  "ParsedResults": [
    {
      "Overlay": {
        "Lines": [
          {
            "LineText": "Juan Jose Carin",
            "Words": [
              {
                "WordText": "Juan",
                "Left": 48,
                "Top": 54.5,
                "Height": 28,
                "Width": 78.5
              },
              {
                "WordText": "Jose",
                "Left": 138.5,
                "Top": 54.5,
                "Height": 28,
                "Width": 74
              },
              {
                "WordText": "Carin",
                "Left": 225.5,
                "Top": 53.5,
                "Height": 29,
                "Width": 88
              }
            ],
            "MaxHeight": 29,
            "MinTop": 53.5
          },
          {
            "LineText": "Data Scientist",
            "Words": [
              {
                "WordText": "Data",
                "Left": 49.5,
                "Top": 98.5,
                "Height": 13.5,
                "Width": 38
              },
              {
                "WordText": "Scientist",
                "Left": 94,
                "Top": 98,
                "Height": 14,
                "Width": 72
              }
            ],
            "MaxHeight": 14,
            "MinTop": 98
          },
          {
            "LineText": "Professional Profile",
            "Words": [
              {
                "WordText": "Professional",
                "Left": 49,
                "Top": 143.5,
                "Height": 13,
                "Width": 93.5
              },
              {
                "WordText": "Profile",
                "Left": 148.5,
                "Top": 144,
                "Height": 12.5,
                "Width": 49.5
              }
            ],
            "MaxHeight": 13,
            "MinTop": 143.5
          },
          {
            "LineText": "Jan. 2016- Mar. 2016",
            "Words": [
              {
                "WordText": "Jan.",
                "Left": 48,
                "Top": 666,
                "Height": 9.5,
                "Width": 22
              },
              {
                "WordText": "2016-",
                "Left": 75,
                "Top": 666,
                "Height": 9.5,
                "Width": 40
              },
              {
                "WordText": "Mar.",
                "Left": 119.5,
                "Top": 666,
                "Height": 9.5,
                "Width": 26.5
              },
              {
                "WordText": "2016",
                "Left": 150.5,
                "Top": 666,
                "Height": 9.5,
                "Width": 28.5
              }
            ],
            "MaxHeight": 9.5,
            "MinTop": 666
          },
          {
            "LineText": "Jun. 2014 -Sep. 2014",
            "Words": [
              {
                "WordText": "Jun.",
                "Left": 48,
                "Top": 772,
                "Height": 9.5,
                "Width": 23
              },
              {
                "WordText": "2014",
                "Left": 75.5,
                "Top": 772,
                "Height": 9.5,
                "Width": 29
              },
              {
                "WordText": "-Sep.",
                "Left": 108,
                "Top": 772,
                "Height": 12,
                "Width": 35.5
              },
              {
                "WordText": "2014",
                "Left": 148.5,
                "Top": 772,
                "Height": 9.5,
                "Width": 28.5
              }
            ],
            "MaxHeight": 12,
            "MinTop": 772
          },
          {
            "LineText": "Mountain View, CA 94041",
            "Words": [
              {
                "WordText": "Mountain",
                "Left": 628,
                "Top": 71,
                "Height": 9.5,
                "Width": 51.5
              },
              {
                "WordText": "View,",
                "Left": 683.5,
                "Top": 71,
                "Height": 11.5,
                "Width": 29
              },
              {
                "WordText": "CA",
                "Left": 717,
                "Top": 71.5,
                "Height": 9,
                "Width": 14
              },
              {
                "WordText": "94041",
                "Left": 734.5,
                "Top": 71.5,
                "Height": 9,
                "Width": 33
              }
            ],
            "MaxHeight": 11.5,
            "MinTop": 71
          },
          {
            "LineText": "650-336-4590 | juanjose.carin@gmail.com",
            "Words": [
              {
                "WordText": "650-336-4590",
                "Left": 531,
                "Top": 88,
                "Height": 8.5,
                "Width": 75
              },
              {
                "WordText": "|",
                "Left": 614.5,
                "Top": 87,
                "Height": 12,
                "Width": 1.5
              },
              {
                "WordText": "juanjose.carin@gmail.com",
                "Left": 624,
                "Top": 87.5,
                "Height": 11.5,
                "Width": 143.5
              }
            ],
            "MaxHeight": 12,
            "MinTop": 87
          },
          {
            "LineText": "linkedin.com/in/juanjosecarin I",
            "Words": [
              {
                "WordText": "linkedin.com/in/juanjosecarin",
                "Left": 444.5,
                "Top": 103,
                "Height": 12.5,
                "Width": 160.5
              },
              {
                "WordText": "I",
                "Left": 614.5,
                "Top": 103,
                "Height": 12.5,
                "Width": 1.5
              }
            ],
            "MaxHeight": 12.5,
            "MinTop": 103
          },
          {
            "LineText": "juanjocarin.github.io",
            "Words": [
              {
                "WordText": "juanjocarin.github.io",
                "Left": 655.5,
                "Top": 103.5,
                "Height": 12,
                "Width": 112
              }
            ],
            "MaxHeight": 12,
            "MinTop": 103.5
          },
          {
            "LineText": "Passionate about data analysis and experiments, mainly focused on user behavior, experience, and engagement, with a solid",
            "Words": [
              {
                "WordText": "Passionate",
                "Left": 49,
                "Top": 175.5,
                "Height": 9,
                "Width": 57
              },
              {
                "WordText": "about",
                "Left": 112.5,
                "Top": 175.5,
                "Height": 9,
                "Width": 31.5
              },
              {
                "WordText": "data",
                "Left": 150.5,
                "Top": 175.5,
                "Height": 9,
                "Width": 23
              },
              {
                "WordText": "analysis",
                "Left": 180.5,
                "Top": 175.5,
                "Height": 11.5,
                "Width": 41.5
              },
              {
                "WordText": "and",
                "Left": 228.5,
                "Top": 175.5,
                "Height": 9,
                "Width": 19.5
              },
              {
                "WordText": "experiments,",
                "Left": 255,
                "Top": 175.5,
                "Height": 11.5,
                "Width": 69.5
              },
              {
                "WordText": "mainly",
                "Left": 332,
                "Top": 175.5,
                "Height": 11.5,
                "Width": 35.5
              },
              {
                "WordText": "focused",
                "Left": 374,
                "Top": 175.5,
                "Height": 9,
                "Width": 41.5
              },
              {
                "WordText": "on",
                "Left": 422.5,
                "Top": 178,
                "Height": 6.5,
                "Width": 12.5
              },
              {
                "WordText": "user",
                "Left": 443,
                "Top": 178,
                "Height": 6.5,
                "Width": 22.5
              },
              {
                "WordText": "behavior,",
                "Left": 472,
                "Top": 175.5,
                "Height": 11,
                "Width": 49.5
              },
              {
                "WordText": "experience,",
                "Left": 529,
                "Top": 175.5,
                "Height": 11.5,
                "Width": 61.5
              },
              {
                "WordText": "and",
                "Left": 597.5,
                "Top": 175.5,
                "Height": 9,
                "Width": 19.5
              },
              {
                "WordText": "engagement,",
                "Left": 624,
                "Top": 176.5,
                "Height": 10.5,
                "Width": 70
              },
              {
                "WordText": "with",
                "Left": 701,
                "Top": 175.5,
                "Height": 9,
                "Width": 23
              },
              {
                "WordText": "a",
                "Left": 731,
                "Top": 178,
                "Height": 6.5,
                "Width": 5.5
              },
              {
                "WordText": "solid",
                "Left": 743.5,
                "Top": 175.5,
                "Height": 9,
                "Width": 24
              }
            ],
            "MaxHeight": 11.5,
            "MinTop": 175.5
          },
          {
            "LineText": "background in data science and statistics, and extensive experience using data insights to drive business growth.",
            "Words": [
              {
                "WordText": "background",
                "Left": 49,
                "Top": 191.5,
                "Height": 12,
                "Width": 62.5
              },
              {
                "WordText": "in",
                "Left": 116,
                "Top": 192,
                "Height": 9,
                "Width": 8.5
              },
              {
                "WordText": "data",
                "Left": 128.5,
                "Top": 191.5,
                "Height": 9.5,
                "Width": 23.5
              },
              {
                "WordText": "science",
                "Left": 155.5,
                "Top": 192,
                "Height": 9,
                "Width": 39
              },
              {
                "WordText": "and",
                "Left": 198.5,
                "Top": 191.5,
                "Height": 9.5,
                "Width": 19.5
              },
              {
                "WordText": "statistics,",
                "Left": 222,
                "Top": 192,
                "Height": 11,
                "Width": 49
              },
              {
                "WordText": "and",
                "Left": 275.5,
                "Top": 191.5,
                "Height": 9.5,
                "Width": 19
              },
              {
                "WordText": "extensive",
                "Left": 299,
                "Top": 192,
                "Height": 9,
                "Width": 50.5
              },
              {
                "WordText": "experience",
                "Left": 353,
                "Top": 192,
                "Height": 11.5,
                "Width": 59
              },
              {
                "WordText": "using",
                "Left": 416,
                "Top": 192,
                "Height": 11.5,
                "Width": 28
              },
              {
                "WordText": "data",
                "Left": 447.5,
                "Top": 191.5,
                "Height": 9.5,
                "Width": 23
              },
              {
                "WordText": "insights",
                "Left": 475,
                "Top": 191.5,
                "Height": 12,
                "Width": 40
              },
              {
                "WordText": "to",
                "Left": 518.5,
                "Top": 192.5,
                "Height": 8.5,
                "Width": 11
              },
              {
                "WordText": "drive",
                "Left": 533.5,
                "Top": 191.5,
                "Height": 9.5,
                "Width": 26.5
              },
              {
                "WordText": "business",
                "Left": 564,
                "Top": 191.5,
                "Height": 9.5,
                "Width": 45.5
              },
              {
                "WordText": "growth.",
                "Left": 613,
                "Top": 191.5,
                "Height": 12,
                "Width": 41
              }
            ],
            "MaxHeight": 12,
            "MinTop": 191.5
          },
          {
            "LineText": "Education",
            "Words": [
              {
                "WordText": "Education",
                "Left": 49,
                "Top": 221.5,
                "Height": 13,
                "Width": 75.5
              }
            ],
            "MaxHeight": 13,
            "MinTop": 221.5
          },
          {
            "LineText": "2016",
            "Words": [
              {
                "WordText": "2016",
                "Left": 48.5,
                "Top": 254.5,
                "Height": 9.5,
                "Width": 29
              }
            ],
            "MaxHeight": 9.5,
            "MinTop": 254.5
          },
          {
            "LineText": "2014",
            "Words": [
              {
                "WordText": "2014",
                "Left": 48.5,
                "Top": 343,
                "Height": 9.5,
                "Width": 29
              }
            ],
            "MaxHeight": 9.5,
            "MinTop": 343
          },
          {
            "LineText": "2005",
            "Words": [
              {
                "WordText": "2005",
                "Left": 48.5,
                "Top": 431.5,
                "Height": 9,
                "Width": 28.5
              }
            ],
            "MaxHeight": 9,
            "MinTop": 431.5
          },
          {
            "LineText": "University of California, Berkeley",
            "Words": [
              {
                "WordText": "University",
                "Left": 97,
                "Top": 253.5,
                "Height": 13,
                "Width": 61
              },
              {
                "WordText": "of",
                "Left": 162,
                "Top": 253.5,
                "Height": 10.5,
                "Width": 12.5
              },
              {
                "WordText": "California,",
                "Left": 177.5,
                "Top": 253.5,
                "Height": 12.5,
                "Width": 61.5
              },
              {
                "WordText": "Berkeley",
                "Left": 244,
                "Top": 253.5,
                "Height": 13,
                "Width": 52
              }
            ],
            "MaxHeight": 13,
            "MinTop": 253.5
          },
          {
            "LineText": "Relevant courses:",
            "Words": [
              {
                "WordText": "Relevant",
                "Left": 96.5,
                "Top": 271.5,
                "Height": 9,
                "Width": 47
              },
              {
                "WordText": "courses:",
                "Left": 146.5,
                "Top": 274,
                "Height": 6.5,
                "Width": 43
              }
            ],
            "MaxHeight": 9,
            "MinTop": 271.5
          },
          {
            "LineText": "• Machine Learning",
            "Words": [
              {
                "WordText": "•",
                "Left": 109,
                "Top": 290.5,
                "Height": 4.5,
                "Width": 4.5
              },
              {
                "WordText": "Machine",
                "Left": 121,
                "Top": 287.5,
                "Height": 9.5,
                "Width": 46
              },
              {
                "WordText": "Learning",
                "Left": 171,
                "Top": 287.5,
                "Height": 12,
                "Width": 45.5
              }
            ],
            "MaxHeight": 12,
            "MinTop": 287.5
          },
          {
            "LineText": "• Machine Learning at Scale",
            "Words": [
              {
                "WordText": "•",
                "Left": 109,
                "Top": 307,
                "Height": 4,
                "Width": 4.5
              },
              {
                "WordText": "Machine",
                "Left": 121,
                "Top": 304,
                "Height": 9,
                "Width": 46
              },
              {
                "WordText": "Learning",
                "Left": 171,
                "Top": 304,
                "Height": 11.5,
                "Width": 45.5
              },
              {
                "WordText": "at",
                "Left": 220,
                "Top": 305,
                "Height": 8,
                "Width": 10.5
              },
              {
                "WordText": "Scale",
                "Left": 234,
                "Top": 304,
                "Height": 9,
                "Width": 27
              }
            ],
            "MaxHeight": 11.5,
            "MinTop": 304
          },
          {
            "LineText": "• Storing and Retrieving Data",
            "Words": [
              {
                "WordText": "•",
                "Left": 109,
                "Top": 323,
                "Height": 4.5,
                "Width": 4.5
              },
              {
                "WordText": "Storing",
                "Left": 120.5,
                "Top": 320.5,
                "Height": 11.5,
                "Width": 38
              },
              {
                "WordText": "and",
                "Left": 162,
                "Top": 320,
                "Height": 9.5,
                "Width": 19.5
              },
              {
                "WordText": "Retrieving",
                "Left": 186,
                "Top": 320.5,
                "Height": 11.5,
                "Width": 54
              },
              {
                "WordText": "Data",
                "Left": 244,
                "Top": 320.5,
                "Height": 9,
                "Width": 24
              }
            ],
            "MaxHeight": 11.5,
            "MinTop": 320
          },
          {
            "LineText": "Universidad Politécnica de Madrid",
            "Words": [
              {
                "WordText": "Universidad",
                "Left": 97,
                "Top": 342,
                "Height": 10.5,
                "Width": 72
              },
              {
                "WordText": "Politécnica",
                "Left": 172.5,
                "Top": 341.5,
                "Height": 11,
                "Width": 65.5
              },
              {
                "WordText": "de",
                "Left": 242,
                "Top": 342,
                "Height": 10.5,
                "Width": 14.5
              },
              {
                "WordText": "Madrid",
                "Left": 260,
                "Top": 342,
                "Height": 10.5,
                "Width": 45
              }
            ],
            "MaxHeight": 11,
            "MinTop": 341.5
          },
          {
            "LineText": "Relevant courses:",
            "Words": [
              {
                "WordText": "Relevant",
                "Left": 96.5,
                "Top": 359.5,
                "Height": 9.5,
                "Width": 47
              },
              {
                "WordText": "courses:",
                "Left": 146.5,
                "Top": 362,
                "Height": 7,
                "Width": 43
              }
            ],
            "MaxHeight": 9.5,
            "MinTop": 359.5
          },
          {
            "LineText": "• Data Mining",
            "Words": [
              {
                "WordText": "•",
                "Left": 109,
                "Top": 379,
                "Height": 4.5,
                "Width": 4.5
              },
              {
                "WordText": "Data",
                "Left": 121,
                "Top": 376.5,
                "Height": 9,
                "Width": 23.5
              },
              {
                "WordText": "Mining",
                "Left": 149.5,
                "Top": 376,
                "Height": 12,
                "Width": 36.5
              }
            ],
            "MaxHeight": 12,
            "MinTop": 376
          },
          {
            "LineText": "• Multivariate Analysis",
            "Words": [
              {
                "WordText": "•",
                "Left": 109,
                "Top": 395,
                "Height": 4.5,
                "Width": 4.5
              },
              {
                "WordText": "Multivariate",
                "Left": 121,
                "Top": 392,
                "Height": 9.5,
                "Width": 65
              },
              {
                "WordText": "Analysis",
                "Left": 189.5,
                "Top": 392,
                "Height": 12,
                "Width": 43.5
              }
            ],
            "MaxHeight": 12,
            "MinTop": 392
          },
          {
            "LineText": "• Time Series",
            "Words": [
              {
                "WordText": "•",
                "Left": 109,
                "Top": 411.5,
                "Height": 4.5,
                "Width": 4.5
              },
              {
                "WordText": "Time",
                "Left": 120,
                "Top": 408.5,
                "Height": 9.5,
                "Width": 26.5
              },
              {
                "WordText": "Series",
                "Left": 150,
                "Top": 408.5,
                "Height": 9.5,
                "Width": 32
              }
            ],
            "MaxHeight": 9.5,
            "MinTop": 408.5
          },
          {
            "LineText": "Universidad Politécnica de Madrid",
            "Words": [
              {
                "WordText": "Universidad",
                "Left": 97,
                "Top": 430.5,
                "Height": 10,
                "Width": 72
              },
              {
                "WordText": "Politécnica",
                "Left": 172.5,
                "Top": 430,
                "Height": 10.5,
                "Width": 65.5
              },
              {
                "WordText": "de",
                "Left": 242,
                "Top": 430.5,
                "Height": 10,
                "Width": 14.5
              },
              {
                "WordText": "Madrid",
                "Left": 260,
                "Top": 430.5,
                "Height": 10,
                "Width": 45
              }
            ],
            "MaxHeight": 10.5,
            "MinTop": 430
          },
          {
            "LineText": "Master of Information and Data Science",
            "Words": [
              {
                "WordText": "Master",
                "Left": 327.5,
                "Top": 254.5,
                "Height": 9.5,
                "Width": 41.5
              },
              {
                "WordText": "of",
                "Left": 373,
                "Top": 253.5,
                "Height": 10.5,
                "Width": 12
              },
              {
                "WordText": "Information",
                "Left": 389,
                "Top": 253.5,
                "Height": 10.5,
                "Width": 69.5
              },
              {
                "WordText": "and",
                "Left": 463,
                "Top": 253.5,
                "Height": 10.5,
                "Width": 21
              },
              {
                "WordText": "Data",
                "Left": 489.5,
                "Top": 254.5,
                "Height": 9.5,
                "Width": 26
              },
              {
                "WordText": "Science",
                "Left": 520,
                "Top": 254.5,
                "Height": 9.5,
                "Width": 44
              }
            ],
            "MaxHeight": 10.5,
            "MinTop": 253.5
          },
          {
            "LineText": "• Field Experiments",
            "Words": [
              {
                "WordText": "•",
                "Left": 333,
                "Top": 274.5,
                "Height": 4,
                "Width": 4.5
              },
              {
                "WordText": "Field",
                "Left": 344.5,
                "Top": 271.5,
                "Height": 9,
                "Width": 24.5
              },
              {
                "WordText": "Experiments",
                "Left": 373.5,
                "Top": 271.5,
                "Height": 11.5,
                "Width": 66.5
              }
            ],
            "MaxHeight": 11.5,
            "MinTop": 271.5
          },
          {
            "LineText": "• Applied Regression and Time Series",
            "Words": [
              {
                "WordText": "•",
                "Left": 333,
                "Top": 290.5,
                "Height": 4.5,
                "Width": 4.5
              },
              {
                "WordText": "Applied",
                "Left": 344,
                "Top": 287.5,
                "Height": 12,
                "Width": 40.5
              },
              {
                "WordText": "Regression",
                "Left": 389,
                "Top": 287.5,
                "Height": 12,
                "Width": 57.5
              },
              {
                "WordText": "and",
                "Left": 450.5,
                "Top": 287.5,
                "Height": 9.5,
                "Width": 19.5
              },
              {
                "WordText": "Time",
                "Left": 473.5,
                "Top": 287.5,
                "Height": 9.5,
                "Width": 26.5
              },
              {
                "WordText": "Series",
                "Left": 504,
                "Top": 287.5,
                "Height": 9.5,
                "Width": 31.5
              }
            ],
            "MaxHeight": 12,
            "MinTop": 287.5
          },
          {
            "LineText": "Analysis",
            "Words": [
              {
                "WordText": "Analysis",
                "Left": 344,
                "Top": 304,
                "Height": 11.5,
                "Width": 43
              }
            ],
            "MaxHeight": 11.5,
            "MinTop": 304
          },
          {
            "LineText": "• Exploring and Analyzing Data",
            "Words": [
              {
                "WordText": "•",
                "Left": 333,
                "Top": 323,
                "Height": 4.5,
                "Width": 4.5
              },
              {
                "WordText": "Exploring",
                "Left": 344.5,
                "Top": 320,
                "Height": 12,
                "Width": 49.5
              },
              {
                "WordText": "and",
                "Left": 397.5,
                "Top": 320,
                "Height": 9.5,
                "Width": 19.5
              },
              {
                "WordText": "Analyzing",
                "Left": 421,
                "Top": 320,
                "Height": 12,
                "Width": 51.5
              },
              {
                "WordText": "Data",
                "Left": 476.5,
                "Top": 320.5,
                "Height": 9,
                "Width": 24
              }
            ],
            "MaxHeight": 12,
            "MinTop": 320
          },
          {
            "LineText": "•",
            "Words": [
              {
                "WordText": "•",
                "Left": 556.5,
                "Top": 274.5,
                "Height": 4,
                "Width": 4.5
              }
            ],
            "MaxHeight": 4,
            "MinTop": 274.5
          },
          {
            "LineText": "•",
            "Words": [
              {
                "WordText": "•",
                "Left": 556.5,
                "Top": 307,
                "Height": 4,
                "Width": 4.5
              }
            ],
            "MaxHeight": 4,
            "MinTop": 307
          },
          {
            "LineText": "GPA: 3.93",
            "Words": [
              {
                "WordText": "GPA:",
                "Left": 715.5,
                "Top": 254,
                "Height": 9,
                "Width": 25
              },
              {
                "WordText": "3.93",
                "Left": 745,
                "Top": 254,
                "Height": 9,
                "Width": 22.5
              }
            ],
            "MaxHeight": 9,
            "MinTop": 254
          },
          {
            "LineText": "Data Visualization and",
            "Words": [
              {
                "WordText": "Data",
                "Left": 568.5,
                "Top": 272,
                "Height": 8.5,
                "Width": 23.5
              },
              {
                "WordText": "Visualization",
                "Left": 596,
                "Top": 271.5,
                "Height": 9,
                "Width": 67.5
              },
              {
                "WordText": "and",
                "Left": 668,
                "Top": 271.5,
                "Height": 9,
                "Width": 19
              }
            ],
            "MaxHeight": 9,
            "MinTop": 271.5
          },
          {
            "LineText": "Communication",
            "Words": [
              {
                "WordText": "Communication",
                "Left": 568,
                "Top": 287.5,
                "Height": 9.5,
                "Width": 84.5
              }
            ],
            "MaxHeight": 9.5,
            "MinTop": 287.5
          },
          {
            "LineText": "Research Design and Applications for",
            "Words": [
              {
                "WordText": "Research",
                "Left": 568.5,
                "Top": 304,
                "Height": 9,
                "Width": 47.5
              },
              {
                "WordText": "Design",
                "Left": 621,
                "Top": 304,
                "Height": 11.5,
                "Width": 34.5
              },
              {
                "WordText": "and",
                "Left": 659.5,
                "Top": 304,
                "Height": 9,
                "Width": 19.5
              },
              {
                "WordText": "Applications",
                "Left": 683,
                "Top": 304,
                "Height": 11.5,
                "Width": 66
              },
              {
                "WordText": "for",
                "Left": 752.5,
                "Top": 304,
                "Height": 9,
                "Width": 15.5
              }
            ],
            "MaxHeight": 11.5,
            "MinTop": 304
          },
          {
            "LineText": "Data Analysis",
            "Words": [
              {
                "WordText": "Data",
                "Left": 568.5,
                "Top": 320.5,
                "Height": 9,
                "Width": 23.5
              },
              {
                "WordText": "Analysis",
                "Left": 596,
                "Top": 320,
                "Height": 12,
                "Width": 43.5
              }
            ],
            "MaxHeight": 12,
            "MinTop": 320
          },
          {
            "LineText": "M.S. in Statistical and Computational Information Processing",
            "Words": [
              {
                "WordText": "M.S.",
                "Left": 327.5,
                "Top": 343,
                "Height": 9.5,
                "Width": 25
              },
              {
                "WordText": "in",
                "Left": 357.5,
                "Top": 343,
                "Height": 9.5,
                "Width": 9
              },
              {
                "WordText": "Statistical",
                "Left": 371.5,
                "Top": 342,
                "Height": 10.5,
                "Width": 56
              },
              {
                "WordText": "and",
                "Left": 432.5,
                "Top": 342,
                "Height": 10.5,
                "Width": 21
              },
              {
                "WordText": "Computational",
                "Left": 458.5,
                "Top": 342,
                "Height": 13,
                "Width": 87
              },
              {
                "WordText": "Information",
                "Left": 551,
                "Top": 342,
                "Height": 10.5,
                "Width": 69
              },
              {
                "WordText": "Processing",
                "Left": 625,
                "Top": 343,
                "Height": 12,
                "Width": 62
              }
            ],
            "MaxHeight": 13,
            "MinTop": 342
          },
          {
            "LineText": "GPA: 3.69",
            "Words": [
              {
                "WordText": "GPA:",
                "Left": 715.5,
                "Top": 342.5,
                "Height": 9,
                "Width": 25
              },
              {
                "WordText": "3.69",
                "Left": 745,
                "Top": 342.5,
                "Height": 9,
                "Width": 22.5
              }
            ],
            "MaxHeight": 9,
            "MinTop": 342.5
          },
          {
            "LineText": "• Neural Networks and Statistical",
            "Words": [
              {
                "WordText": "•",
                "Left": 333,
                "Top": 362.5,
                "Height": 4.5,
                "Width": 4.5
              },
              {
                "WordText": "Neural",
                "Left": 344.5,
                "Top": 359.5,
                "Height": 9.5,
                "Width": 35
              },
              {
                "WordText": "Networks",
                "Left": 384,
                "Top": 359.5,
                "Height": 9.5,
                "Width": 51
              },
              {
                "WordText": "and",
                "Left": 438.5,
                "Top": 359.5,
                "Height": 9.5,
                "Width": 19.5
              },
              {
                "WordText": "Statistical",
                "Left": 462,
                "Top": 359.5,
                "Height": 9.5,
                "Width": 51
              }
            ],
            "MaxHeight": 9.5,
            "MinTop": 359.5
          },
          {
            "LineText": "Learning",
            "Words": [
              {
                "WordText": "Learning",
                "Left": 344.5,
                "Top": 376,
                "Height": 12,
                "Width": 45.5
              }
            ],
            "MaxHeight": 12,
            "MinTop": 376
          },
          {
            "LineText": "• Monte Carlo Techniques",
            "Words": [
              {
                "WordText": "•",
                "Left": 556.5,
                "Top": 362.5,
                "Height": 4.5,
                "Width": 4.5
              },
              {
                "WordText": "Monte",
                "Left": 568.5,
                "Top": 360,
                "Height": 9,
                "Width": 35
              },
              {
                "WordText": "Carlo",
                "Left": 607.5,
                "Top": 359.5,
                "Height": 9.5,
                "Width": 27.5
              },
              {
                "WordText": "Techniques",
                "Left": 638,
                "Top": 359.5,
                "Height": 12,
                "Width": 61.5
              }
            ],
            "MaxHeight": 12,
            "MinTop": 359.5
          },
          {
            "LineText": "• Numerical Methods in Finance",
            "Words": [
              {
                "WordText": "•",
                "Left": 556.5,
                "Top": 379,
                "Height": 4.5,
                "Width": 4.5
              },
              {
                "WordText": "Numerical",
                "Left": 568.5,
                "Top": 376,
                "Height": 9.5,
                "Width": 53.5
              },
              {
                "WordText": "Methods",
                "Left": 627,
                "Top": 376,
                "Height": 9,
                "Width": 47.5
              },
              {
                "WordText": "in",
                "Left": 678.5,
                "Top": 376,
                "Height": 9.5,
                "Width": 8.5
              },
              {
                "WordText": "Finance",
                "Left": 692,
                "Top": 376,
                "Height": 9.5,
                "Width": 40.5
              }
            ],
            "MaxHeight": 9.5,
            "MinTop": 376
          },
          {
            "LineText": "• Regression and Prediction Methods",
            "Words": [
              {
                "WordText": "•",
                "Left": 333,
                "Top": 395,
                "Height": 4.5,
                "Width": 4.5
              },
              {
                "WordText": "Regression",
                "Left": 344.5,
                "Top": 392.5,
                "Height": 11.5,
                "Width": 57.5
              },
              {
                "WordText": "and",
                "Left": 406,
                "Top": 392,
                "Height": 9.5,
                "Width": 19.5
              },
              {
                "WordText": "Prediction",
                "Left": 430,
                "Top": 392,
                "Height": 9.5,
                "Width": 54
              },
              {
                "WordText": "Methods",
                "Left": 488.5,
                "Top": 392,
                "Height": 9.5,
                "Width": 47.5
              }
            ],
            "MaxHeight": 11.5,
            "MinTop": 392
          },
          {
            "LineText": "• Stochastic Models in Finance",
            "Words": [
              {
                "WordText": "•",
                "Left": 556.5,
                "Top": 395,
                "Height": 4.5,
                "Width": 4.5
              },
              {
                "WordText": "Stochastic",
                "Left": 568,
                "Top": 392,
                "Height": 9.5,
                "Width": 54
              },
              {
                "WordText": "Models",
                "Left": 626.5,
                "Top": 392,
                "Height": 9.5,
                "Width": 39
              },
              {
                "WordText": "in",
                "Left": 669.5,
                "Top": 392.5,
                "Height": 9,
                "Width": 8.5
              },
              {
                "WordText": "Finance",
                "Left": 683,
                "Top": 392.5,
                "Height": 9,
                "Width": 40.5
              }
            ],
            "MaxHeight": 9.5,
            "MinTop": 392
          },
          {
            "LineText": "• Optimization Techniques",
            "Words": [
              {
                "WordText": "•",
                "Left": 333,
                "Top": 411.5,
                "Height": 4.5,
                "Width": 4.5
              },
              {
                "WordText": "Optimization",
                "Left": 344,
                "Top": 408.5,
                "Height": 12,
                "Width": 69
              },
              {
                "WordText": "Techniques",
                "Left": 417,
                "Top": 408.5,
                "Height": 12,
                "Width": 61.5
              }
            ],
            "MaxHeight": 12,
            "MinTop": 408.5
          },
          {
            "LineText": "• Bayesian Networks",
            "Words": [
              {
                "WordText": "•",
                "Left": 556.5,
                "Top": 411.5,
                "Height": 4.5,
                "Width": 4.5
              },
              {
                "WordText": "Bayesian",
                "Left": 568.5,
                "Top": 408.5,
                "Height": 12,
                "Width": 46
              },
              {
                "WordText": "Networks",
                "Left": 619.5,
                "Top": 408.5,
                "Height": 9.5,
                "Width": 51
              }
            ],
            "MaxHeight": 12,
            "MinTop": 408.5
          },
          {
            "LineText": "M.S. in Telecommunication Engineering",
            "Words": [
              {
                "WordText": "M.S.",
                "Left": 327.5,
                "Top": 431.5,
                "Height": 9,
                "Width": 25
              },
              {
                "WordText": "in",
                "Left": 357.5,
                "Top": 431.5,
                "Height": 9,
                "Width": 9
              },
              {
                "WordText": "Telecommunication",
                "Left": 371,
                "Top": 430.5,
                "Height": 10,
                "Width": 117.5
              },
              {
                "WordText": "Engineering",
                "Left": 493.5,
                "Top": 431.5,
                "Height": 11.5,
                "Width": 69
              }
            ],
            "MaxHeight": 11.5,
            "MinTop": 430.5
          },
          {
            "LineText": "GPA: 3.03",
            "Words": [
              {
                "WordText": "GPA:",
                "Left": 715.5,
                "Top": 430.5,
                "Height": 9,
                "Width": 25
              },
              {
                "WordText": "3.03",
                "Left": 745,
                "Top": 430.5,
                "Height": 9,
                "Width": 22.5
              }
            ],
            "MaxHeight": 9,
            "MinTop": 430.5
          },
          {
            "LineText": "Focus Area:",
            "Words": [
              {
                "WordText": "Focus",
                "Left": 96.5,
                "Top": 448.5,
                "Height": 9,
                "Width": 30
              },
              {
                "WordText": "Area:",
                "Left": 129,
                "Top": 448.5,
                "Height": 9,
                "Width": 29.5
              }
            ],
            "MaxHeight": 9,
            "MinTop": 448.5
          },
          {
            "LineText": "Fellowship:",
            "Words": [
              {
                "WordText": "Fellowship:",
                "Left": 96.5,
                "Top": 464.5,
                "Height": 11.5,
                "Width": 59.5
              }
            ],
            "MaxHeight": 11.5,
            "MinTop": 464.5
          },
          {
            "LineText": "Radio communication systems (radar and mobile).",
            "Words": [
              {
                "WordText": "Radio",
                "Left": 193,
                "Top": 448,
                "Height": 9.5,
                "Width": 29.5
              },
              {
                "WordText": "communication",
                "Left": 226,
                "Top": 448,
                "Height": 9.5,
                "Width": 83.5
              },
              {
                "WordText": "systems",
                "Left": 314,
                "Top": 449,
                "Height": 11,
                "Width": 42.5
              },
              {
                "WordText": "(radar",
                "Left": 360.5,
                "Top": 447.5,
                "Height": 12.5,
                "Width": 32.5
              },
              {
                "WordText": "and",
                "Left": 396.5,
                "Top": 448,
                "Height": 9.5,
                "Width": 19.5
              },
              {
                "WordText": "mobile).",
                "Left": 420.5,
                "Top": 447.5,
                "Height": 12.5,
                "Width": 43
              }
            ],
            "MaxHeight": 12.5,
            "MinTop": 447.5
          },
          {
            "LineText": "First year at University, due to Honors obtained last year at high school.",
            "Words": [
              {
                "WordText": "First",
                "Left": 193,
                "Top": 464.5,
                "Height": 9,
                "Width": 22
              },
              {
                "WordText": "year",
                "Left": 218.5,
                "Top": 467,
                "Height": 9,
                "Width": 23.5
              },
              {
                "WordText": "at",
                "Left": 245.5,
                "Top": 465.5,
                "Height": 8,
                "Width": 10.5
              },
              {
                "WordText": "University,",
                "Left": 260,
                "Top": 464.5,
                "Height": 11.5,
                "Width": 56.5
              },
              {
                "WordText": "due",
                "Left": 320.5,
                "Top": 464.5,
                "Height": 9,
                "Width": 20
              },
              {
                "WordText": "to",
                "Left": 344,
                "Top": 465.5,
                "Height": 8,
                "Width": 11
              },
              {
                "WordText": "Honors",
                "Left": 359.5,
                "Top": 465,
                "Height": 8.5,
                "Width": 38
              },
              {
                "WordText": "obtained",
                "Left": 401,
                "Top": 464.5,
                "Height": 9,
                "Width": 47.5
              },
              {
                "WordText": "last",
                "Left": 453,
                "Top": 464.5,
                "Height": 9,
                "Width": 18.5
              },
              {
                "WordText": "year",
                "Left": 474.5,
                "Top": 467,
                "Height": 9,
                "Width": 23.5
              },
              {
                "WordText": "at",
                "Left": 501.5,
                "Top": 465.5,
                "Height": 8,
                "Width": 10.5
              },
              {
                "WordText": "high",
                "Left": 516,
                "Top": 464.5,
                "Height": 11.5,
                "Width": 22
              },
              {
                "WordText": "school.",
                "Left": 542,
                "Top": 464.5,
                "Height": 9,
                "Width": 37
              }
            ],
            "MaxHeight": 11.5,
            "MinTop": 464.5
          },
          {
            "LineText": "Skills",
            "Words": [
              {
                "WordText": "Skills",
                "Left": 48.5,
                "Top": 494.5,
                "Height": 12.5,
                "Width": 38
              }
            ],
            "MaxHeight": 12.5,
            "MinTop": 494.5
          },
          {
            "LineText": "Proficient:",
            "Words": [
              {
                "WordText": "Proficient:",
                "Left": 49,
                "Top": 543.5,
                "Height": 9.5,
                "Width": 54
              }
            ],
            "MaxHeight": 9.5,
            "MinTop": 543.5
          },
          {
            "LineText": "Intermediate:",
            "Words": [
              {
                "WordText": "Intermediate:",
                "Left": 49,
                "Top": 561.5,
                "Height": 9,
                "Width": 72.5
              }
            ],
            "MaxHeight": 9,
            "MinTop": 561.5
          },
          {
            "LineText": "Basic:",
            "Words": [
              {
                "WordText": "Basic:",
                "Left": 49,
                "Top": 579.5,
                "Height": 9,
                "Width": 29.5
              }
            ],
            "MaxHeight": 9,
            "MinTop": 579.5
          },
          {
            "LineText": "Experience",
            "Words": [
              {
                "WordText": "Experience",
                "Left": 49,
                "Top": 610.5,
                "Height": 15,
                "Width": 83.5
              }
            ],
            "MaxHeight": 15,
            "MinTop": 610.5
          },
          {
            "LineText": "DATA SCIENCE",
            "Words": [
              {
                "WordText": "DATA",
                "Left": 49,
                "Top": 642.5,
                "Height": 9.5,
                "Width": 32
              },
              {
                "WordText": "SCIENCE",
                "Left": 85,
                "Top": 642.5,
                "Height": 9.5,
                "Width": 49
              }
            ],
            "MaxHeight": 9.5,
            "MinTop": 642.5
          },
          {
            "LineText": "Programming / Statistics",
            "Words": [
              {
                "WordText": "Programming",
                "Left": 145,
                "Top": 526.5,
                "Height": 11.5,
                "Width": 72.5
              },
              {
                "WordText": "/",
                "Left": 221,
                "Top": 525.5,
                "Height": 10.5,
                "Width": 4.5
              },
              {
                "WordText": "Statistics",
                "Left": 229,
                "Top": 526.5,
                "Height": 9,
                "Width": 47.5
              }
            ],
            "MaxHeight": 11.5,
            "MinTop": 525.5
          },
          {
            "LineText": "R, Python, SQL",
            "Words": [
              {
                "WordText": "R,",
                "Left": 144.5,
                "Top": 543.5,
                "Height": 11.5,
                "Width": 10.5
              },
              {
                "WordText": "Python,",
                "Left": 159.5,
                "Top": 542.5,
                "Height": 13,
                "Width": 44
              },
              {
                "WordText": "SQL",
                "Left": 207.5,
                "Top": 543.5,
                "Height": 10,
                "Width": 22
              }
            ],
            "MaxHeight": 13,
            "MinTop": 542.5
          },
          {
            "LineText": "spss, SAS, Matiab",
            "Words": [
              {
                "WordText": "spss,",
                "Left": 144,
                "Top": 561.5,
                "Height": 11,
                "Width": 30.5
              },
              {
                "WordText": "SAS,",
                "Left": 178.5,
                "Top": 561.5,
                "Height": 11,
                "Width": 24.5
              },
              {
                "WordText": "Matiab",
                "Left": 207.5,
                "Top": 560.5,
                "Height": 10,
                "Width": 43
              }
            ],
            "MaxHeight": 11,
            "MinTop": 560.5
          },
          {
            "LineText": "E Views, Demetra+",
            "Words": [
              {
                "WordText": "E",
                "Left": 144.5,
                "Top": 579,
                "Height": 9.5,
                "Width": 7
              },
              {
                "WordText": "Views,",
                "Left": 152.5,
                "Top": 579,
                "Height": 11.5,
                "Width": 36.5
              },
              {
                "WordText": "Demetra+",
                "Left": 193.5,
                "Top": 579,
                "Height": 9.5,
                "Width": 59
              }
            ],
            "MaxHeight": 11.5,
            "MinTop": 579
          },
          {
            "LineText": "Data Scientist",
            "Words": [
              {
                "WordText": "Data",
                "Left": 205,
                "Top": 666,
                "Height": 9,
                "Width": 27
              },
              {
                "WordText": "Scientist",
                "Left": 236.5,
                "Top": 665,
                "Height": 10,
                "Width": 51
              }
            ],
            "MaxHeight": 10,
            "MinTop": 665
          },
          {
            "LineText": "CONENTO",
            "Words": [
              {
                "WordText": "CONENTO",
                "Left": 204.5,
                "Top": 683.5,
                "Height": 9.5,
                "Width": 59.5
              }
            ],
            "MaxHeight": 9.5,
            "MinTop": 683.5
          },
          {
            "LineText": "Big Data",
            "Words": [
              {
                "WordText": "Big",
                "Left": 301,
                "Top": 526.5,
                "Height": 11.5,
                "Width": 15.5
              },
              {
                "WordText": "Data",
                "Left": 320.5,
                "Top": 526.5,
                "Height": 9,
                "Width": 24
              }
            ],
            "MaxHeight": 11.5,
            "MinTop": 526.5
          },
          {
            "LineText": "Hadoop, Hive, MrJob",
            "Words": [
              {
                "WordText": "Hadoop,",
                "Left": 300.5,
                "Top": 542.5,
                "Height": 13,
                "Width": 49
              },
              {
                "WordText": "Hive,",
                "Left": 354.5,
                "Top": 543.5,
                "Height": 11.5,
                "Width": 28
              },
              {
                "WordText": "MrJob",
                "Left": 387.5,
                "Top": 542.5,
                "Height": 10.5,
                "Width": 36.5
              }
            ],
            "MaxHeight": 13,
            "MinTop": 542.5
          },
          {
            "LineText": "Spark, Storm",
            "Words": [
              {
                "WordText": "Spark,",
                "Left": 300,
                "Top": 560.5,
                "Height": 12.5,
                "Width": 36.5
              },
              {
                "WordText": "Storm",
                "Left": 340.5,
                "Top": 561.5,
                "Height": 9,
                "Width": 35
              }
            ],
            "MaxHeight": 12.5,
            "MinTop": 560.5
          },
          {
            "LineText": "Visualization",
            "Words": [
              {
                "WordText": "Visualization",
                "Left": 456.5,
                "Top": 526,
                "Height": 9.5,
                "Width": 67.5
              }
            ],
            "MaxHeight": 9.5,
            "MinTop": 526
          },
          {
            "LineText": "Tableau",
            "Words": [
              {
                "WordText": "Tableau",
                "Left": 457,
                "Top": 542.5,
                "Height": 10.5,
                "Width": 46.5
              }
            ],
            "MaxHeight": 10.5,
            "MinTop": 542.5
          },
          {
            "LineText": "D3.js",
            "Words": [
              {
                "WordText": "D3.js",
                "Left": 456.5,
                "Top": 579,
                "Height": 12,
                "Width": 29
              }
            ],
            "MaxHeight": 12,
            "MinTop": 579
          },
          {
            "LineText": "Others",
            "Words": [
              {
                "WordText": "Others",
                "Left": 612.5,
                "Top": 526,
                "Height": 9.5,
                "Width": 36
              }
            ],
            "MaxHeight": 9.5,
            "MinTop": 526
          },
          {
            "LineText": "Git, A WS",
            "Words": [
              {
                "WordText": "Git,",
                "Left": 612.5,
                "Top": 543.5,
                "Height": 11.5,
                "Width": 20
              },
              {
                "WordText": "A",
                "Left": 636.5,
                "Top": 543.5,
                "Height": 9.5,
                "Width": 8
              },
              {
                "WordText": "WS",
                "Left": 646.5,
                "Top": 543.5,
                "Height": 9.5,
                "Width": 18.5
              }
            ],
            "MaxHeight": 11.5,
            "MinTop": 543.5
          },
          {
            "LineText": "Bash",
            "Words": [
              {
                "WordText": "Bash",
                "Left": 612.5,
                "Top": 560.5,
                "Height": 10,
                "Width": 27.5
              }
            ],
            "MaxHeight": 10,
            "MinTop": 560.5
          },
          {
            "LineText": "Gephi, Ne04j, QGIS",
            "Words": [
              {
                "WordText": "Gephi,",
                "Left": 613,
                "Top": 578,
                "Height": 13,
                "Width": 36.5
              },
              {
                "WordText": "Ne04j,",
                "Left": 654.5,
                "Top": 579,
                "Height": 12,
                "Width": 36.5
              },
              {
                "WordText": "QGIS",
                "Left": 696.5,
                "Top": 579,
                "Height": 10.5,
                "Width": 28.5
              }
            ],
            "MaxHeight": 13,
            "MinTop": 578
          },
          {
            "LineText": "Madrid, Spain (working remotely)",
            "Words": [
              {
                "WordText": "Madrid,",
                "Left": 586.5,
                "Top": 682.5,
                "Height": 11.5,
                "Width": 41
              },
              {
                "WordText": "Spain",
                "Left": 632,
                "Top": 683,
                "Height": 11.5,
                "Width": 28.5
              },
              {
                "WordText": "(working",
                "Left": 665,
                "Top": 682,
                "Height": 12.5,
                "Width": 46.5
              },
              {
                "WordText": "remotely)",
                "Left": 715.5,
                "Top": 682,
                "Height": 12.5,
                "Width": 52
              }
            ],
            "MaxHeight": 12.5,
            "MinTop": 682
          },
          {
            "LineText": "• Designed and implemented the ETL pipeline for a predictive model of traffic on the main roads in",
            "Words": [
              {
                "WordText": "•",
                "Left": 217,
                "Top": 703.5,
                "Height": 4.5,
                "Width": 4.5
              },
              {
                "WordText": "Designed",
                "Left": 229,
                "Top": 700.5,
                "Height": 12,
                "Width": 48
              },
              {
                "WordText": "and",
                "Left": 282.5,
                "Top": 700.5,
                "Height": 9.5,
                "Width": 19
              },
              {
                "WordText": "implemented",
                "Left": 307,
                "Top": 700.5,
                "Height": 12,
                "Width": 71.5
              },
              {
                "WordText": "the",
                "Left": 383,
                "Top": 700.5,
                "Height": 9.5,
                "Width": 18
              },
              {
                "WordText": "ETL",
                "Left": 406,
                "Top": 701,
                "Height": 9,
                "Width": 18
              },
              {
                "WordText": "pipeline",
                "Left": 428.5,
                "Top": 700.5,
                "Height": 12,
                "Width": 42
              },
              {
                "WordText": "for",
                "Left": 475,
                "Top": 700.5,
                "Height": 9.5,
                "Width": 16
              },
              {
                "WordText": "a",
                "Left": 495.5,
                "Top": 703,
                "Height": 7,
                "Width": 5
              },
              {
                "WordText": "predictive",
                "Left": 506,
                "Top": 700.5,
                "Height": 12,
                "Width": 52.5
              },
              {
                "WordText": "model",
                "Left": 564,
                "Top": 700.5,
                "Height": 9.5,
                "Width": 33
              },
              {
                "WordText": "of",
                "Left": 602,
                "Top": 700.5,
                "Height": 9.5,
                "Width": 11
              },
              {
                "WordText": "traffic",
                "Left": 616.5,
                "Top": 700.5,
                "Height": 9.5,
                "Width": 32
              },
              {
                "WordText": "on",
                "Left": 653,
                "Top": 703,
                "Height": 7,
                "Width": 13
              },
              {
                "WordText": "the",
                "Left": 671,
                "Top": 700.5,
                "Height": 9.5,
                "Width": 17.5
              },
              {
                "WordText": "main",
                "Left": 693.5,
                "Top": 701,
                "Height": 9,
                "Width": 25.5
              },
              {
                "WordText": "roads",
                "Left": 725,
                "Top": 700.5,
                "Height": 9.5,
                "Width": 29
              },
              {
                "WordText": "in",
                "Left": 759,
                "Top": 701,
                "Height": 9,
                "Width": 8.5
              }
            ],
            "MaxHeight": 12,
            "MinTop": 700.5
          },
          {
            "LineText": "eastern Spain (a project for the Spanish government).",
            "Words": [
              {
                "WordText": "eastern",
                "Left": 228.5,
                "Top": 718,
                "Height": 8,
                "Width": 39.5
              },
              {
                "WordText": "Spain",
                "Left": 272,
                "Top": 717,
                "Height": 11.5,
                "Width": 28.5
              },
              {
                "WordText": "(a",
                "Left": 305,
                "Top": 716.5,
                "Height": 12,
                "Width": 9
              },
              {
                "WordText": "project",
                "Left": 319,
                "Top": 717,
                "Height": 11.5,
                "Width": 37.5
              },
              {
                "WordText": "for",
                "Left": 360,
                "Top": 717,
                "Height": 9,
                "Width": 15.5
              },
              {
                "WordText": "the",
                "Left": 378.5,
                "Top": 717,
                "Height": 9,
                "Width": 18
              },
              {
                "WordText": "Spanish",
                "Left": 400,
                "Top": 717,
                "Height": 11.5,
                "Width": 40.5
              },
              {
                "WordText": "government).",
                "Left": 444.5,
                "Top": 716.5,
                "Height": 12,
                "Width": 72.5
              }
            ],
            "MaxHeight": 12,
            "MinTop": 716.5
          },
          {
            "LineText": "• Automated scripts in R to extract, transform, clean (incl. anomaly detection), and load into MySQL",
            "Words": [
              {
                "WordText": "•",
                "Left": 217,
                "Top": 736,
                "Height": 4.5,
                "Width": 4.5
              },
              {
                "WordText": "Automated",
                "Left": 228,
                "Top": 733,
                "Height": 9.5,
                "Width": 60.5
              },
              {
                "WordText": "scripts",
                "Left": 293.5,
                "Top": 733.5,
                "Height": 11.5,
                "Width": 34.5
              },
              {
                "WordText": "in",
                "Left": 332.5,
                "Top": 733.5,
                "Height": 9,
                "Width": 9
              },
              {
                "WordText": "R",
                "Left": 346,
                "Top": 733.5,
                "Height": 9,
                "Width": 7
              },
              {
                "WordText": "to",
                "Left": 356.5,
                "Top": 734,
                "Height": 8.5,
                "Width": 11
              },
              {
                "WordText": "extract,",
                "Left": 372.5,
                "Top": 734,
                "Height": 10.5,
                "Width": 40
              },
              {
                "WordText": "transform,",
                "Left": 417,
                "Top": 733,
                "Height": 11.5,
                "Width": 56.5
              },
              {
                "WordText": "clean",
                "Left": 478.5,
                "Top": 733,
                "Height": 9.5,
                "Width": 27.5
              },
              {
                "WordText": "(incl.",
                "Left": 511,
                "Top": 732.5,
                "Height": 12.5,
                "Width": 25
              },
              {
                "WordText": "anomaly",
                "Left": 541,
                "Top": 733,
                "Height": 12,
                "Width": 46
              },
              {
                "WordText": "detection),",
                "Left": 591,
                "Top": 732.5,
                "Height": 12.5,
                "Width": 58
              },
              {
                "WordText": "and",
                "Left": 654,
                "Top": 733,
                "Height": 9.5,
                "Width": 19.5
              },
              {
                "WordText": "load",
                "Left": 678.5,
                "Top": 733,
                "Height": 9.5,
                "Width": 22
              },
              {
                "WordText": "into",
                "Left": 705.5,
                "Top": 733.5,
                "Height": 9,
                "Width": 21
              },
              {
                "WordText": "MySQL",
                "Left": 731,
                "Top": 733.5,
                "Height": 11.5,
                "Width": 36.5
              }
            ],
            "MaxHeight": 12.5,
            "MinTop": 732.5
          },
          {
            "LineText": "data from multiple data sources: road traffic sensors, accidents, road works, weather.",
            "Words": [
              {
                "WordText": "data",
                "Left": 228.5,
                "Top": 749.5,
                "Height": 9,
                "Width": 23
              },
              {
                "WordText": "from",
                "Left": 255.5,
                "Top": 749.5,
                "Height": 9,
                "Width": 25.5
              },
              {
                "WordText": "multiple",
                "Left": 285.5,
                "Top": 749.5,
                "Height": 11.5,
                "Width": 43.5
              },
              {
                "WordText": "data",
                "Left": 333,
                "Top": 749.5,
                "Height": 9,
                "Width": 23
              },
              {
                "WordText": "sources:",
                "Left": 360.5,
                "Top": 752,
                "Height": 6.5,
                "Width": 43.5
              },
              {
                "WordText": "road",
                "Left": 408.5,
                "Top": 749.5,
                "Height": 9,
                "Width": 23.5
              },
              {
                "WordText": "traffic",
                "Left": 436,
                "Top": 749.5,
                "Height": 9,
                "Width": 32
              },
              {
                "WordText": "sensors,",
                "Left": 471.5,
                "Top": 752,
                "Height": 8.5,
                "Width": 43.5
              },
              {
                "WordText": "accidents,",
                "Left": 519,
                "Top": 749.5,
                "Height": 11,
                "Width": 53.5
              },
              {
                "WordText": "road",
                "Left": 577,
                "Top": 749.5,
                "Height": 9,
                "Width": 23.5
              },
              {
                "WordText": "works,",
                "Left": 604.5,
                "Top": 749.5,
                "Height": 11,
                "Width": 34.5
              },
              {
                "WordText": "weather.",
                "Left": 643,
                "Top": 749.5,
                "Height": 9,
                "Width": 47.5
              }
            ],
            "MaxHeight": 11.5,
            "MinTop": 749.5
          },
          {
            "LineText": "Data Scientist",
            "Words": [
              {
                "WordText": "Data",
                "Left": 205,
                "Top": 772,
                "Height": 9.5,
                "Width": 27
              },
              {
                "WordText": "Scientist",
                "Left": 236.5,
                "Top": 771,
                "Height": 10.5,
                "Width": 51
              }
            ],
            "MaxHeight": 10.5,
            "MinTop": 771
          },
          {
            "LineText": "CONENTO",
            "Words": [
              {
                "WordText": "CONENTO",
                "Left": 204.5,
                "Top": 790,
                "Height": 9.5,
                "Width": 59.5
              }
            ],
            "MaxHeight": 9.5,
            "MinTop": 790
          },
          {
            "LineText": "Madrid, Spain",
            "Words": [
              {
                "WordText": "Madrid,",
                "Left": 693.5,
                "Top": 789,
                "Height": 11,
                "Width": 41.5
              },
              {
                "WordText": "Spain",
                "Left": 739,
                "Top": 789,
                "Height": 11.5,
                "Width": 28.5
              }
            ],
            "MaxHeight": 11.5,
            "MinTop": 789
          },
          {
            "LineText": "• Designed an experiment for Google Spain (conducted in October 2014) to measure the impact of",
            "Words": [
              {
                "WordText": "•",
                "Left": 217,
                "Top": 810,
                "Height": 4,
                "Width": 4.5
              },
              {
                "WordText": "Designed",
                "Left": 229,
                "Top": 807,
                "Height": 11.5,
                "Width": 48
              },
              {
                "WordText": "an",
                "Left": 282.5,
                "Top": 809.5,
                "Height": 6.5,
                "Width": 12.5
              },
              {
                "WordText": "experiment",
                "Left": 300,
                "Top": 807,
                "Height": 11.5,
                "Width": 62
              },
              {
                "WordText": "for",
                "Left": 366.5,
                "Top": 807,
                "Height": 9,
                "Width": 15.5
              },
              {
                "WordText": "Google",
                "Left": 387,
                "Top": 807,
                "Height": 11.5,
                "Width": 37.5
              },
              {
                "WordText": "Spain",
                "Left": 429.5,
                "Top": 807,
                "Height": 11.5,
                "Width": 28.5
              },
              {
                "WordText": "(conducted",
                "Left": 463.5,
                "Top": 806.5,
                "Height": 12,
                "Width": 60
              },
              {
                "WordText": "in",
                "Left": 529,
                "Top": 807,
                "Height": 9,
                "Width": 9
              },
              {
                "WordText": "October",
                "Left": 543.5,
                "Top": 807,
                "Height": 9,
                "Width": 43.5
              },
              {
                "WordText": "2014)",
                "Left": 591.5,
                "Top": 806.5,
                "Height": 12,
                "Width": 30
              },
              {
                "WordText": "to",
                "Left": 626.5,
                "Top": 808,
                "Height": 8,
                "Width": 11
              },
              {
                "WordText": "measure",
                "Left": 643,
                "Top": 809.5,
                "Height": 6.5,
                "Width": 46
              },
              {
                "WordText": "the",
                "Left": 693.5,
                "Top": 807,
                "Height": 9,
                "Width": 17.5
              },
              {
                "WordText": "impact",
                "Left": 716.5,
                "Top": 807,
                "Height": 11.5,
                "Width": 36
              },
              {
                "WordText": "of",
                "Left": 757.5,
                "Top": 807,
                "Height": 9,
                "Width": 11
              }
            ],
            "MaxHeight": 12,
            "MinTop": 806.5
          },
          {
            "LineText": "YouTube ads on the sales of a car manufacturer's dealer network.",
            "Words": [
              {
                "WordText": "YouTube",
                "Left": 228,
                "Top": 823,
                "Height": 9.5,
                "Width": 47
              },
              {
                "WordText": "ads",
                "Left": 279,
                "Top": 823,
                "Height": 9.5,
                "Width": 18
              },
              {
                "WordText": "on",
                "Left": 300.5,
                "Top": 825.5,
                "Height": 7,
                "Width": 13
              },
              {
                "WordText": "the",
                "Left": 317.5,
                "Top": 823,
                "Height": 9.5,
                "Width": 17.5
              },
              {
                "WordText": "sales",
                "Left": 339,
                "Top": 823,
                "Height": 9.5,
                "Width": 25.5
              },
              {
                "WordText": "of",
                "Left": 368.5,
                "Top": 823,
                "Height": 9.5,
                "Width": 11
              },
              {
                "WordText": "a",
                "Left": 382.5,
                "Top": 825.5,
                "Height": 7,
                "Width": 5
              },
              {
                "WordText": "car",
                "Left": 392,
                "Top": 825.5,
                "Height": 7,
                "Width": 16.5
              },
              {
                "WordText": "manufacturer's",
                "Left": 412,
                "Top": 823,
                "Height": 9.5,
                "Width": 81.5
              },
              {
                "WordText": "dealer",
                "Left": 497.5,
                "Top": 823,
                "Height": 9.5,
                "Width": 34
              },
              {
                "WordText": "network.",
                "Left": 535,
                "Top": 823,
                "Height": 9.5,
                "Width": 47.5
              }
            ],
            "MaxHeight": 9.5,
            "MinTop": 823
          },
          {
            "LineText": "• A matched-pair, cluster-randomized design, which involved selecting the test and control groups",
            "Words": [
              {
                "WordText": "•",
                "Left": 217,
                "Top": 842.5,
                "Height": 4.5,
                "Width": 4.5
              },
              {
                "WordText": "A",
                "Left": 228,
                "Top": 840,
                "Height": 9,
                "Width": 7.5
              },
              {
                "WordText": "matched-pair,",
                "Left": 241,
                "Top": 839.5,
                "Height": 12,
                "Width": 75
              },
              {
                "WordText": "cluster-randomized",
                "Left": 321.5,
                "Top": 839.5,
                "Height": 9.5,
                "Width": 104.5
              },
              {
                "WordText": "design,",
                "Left": 431.5,
                "Top": 839.5,
                "Height": 12,
                "Width": 37.5
              },
              {
                "WordText": "which",
                "Left": 475,
                "Top": 839.5,
                "Height": 9.5,
                "Width": 31
              },
              {
                "WordText": "involved",
                "Left": 512,
                "Top": 839.5,
                "Height": 9.5,
                "Width": 44.5
              },
              {
                "WordText": "selecting",
                "Left": 562.5,
                "Top": 839.5,
                "Height": 12,
                "Width": 47
              },
              {
                "WordText": "the",
                "Left": 614.5,
                "Top": 839.5,
                "Height": 9.5,
                "Width": 17.5
              },
              {
                "WordText": "test",
                "Left": 637,
                "Top": 840.5,
                "Height": 8.5,
                "Width": 20.5
              },
              {
                "WordText": "and",
                "Left": 663,
                "Top": 839.5,
                "Height": 9.5,
                "Width": 19
              },
              {
                "WordText": "control",
                "Left": 688,
                "Top": 839.5,
                "Height": 9.5,
                "Width": 37.5
              },
              {
                "WordText": "groups",
                "Left": 731,
                "Top": 842,
                "Height": 9.5,
                "Width": 37
              }
            ],
            "MaxHeight": 12,
            "MinTop": 839.5
          },
          {
            "LineText": "from a sample of 50+ cities in Spain (where geo-targeted ads were possible) based on their sales-",
            "Words": [
              {
                "WordText": "from",
                "Left": 228,
                "Top": 855.5,
                "Height": 9.5,
                "Width": 26
              },
              {
                "WordText": "a",
                "Left": 258.5,
                "Top": 858,
                "Height": 7,
                "Width": 5.5
              },
              {
                "WordText": "sample",
                "Left": 269,
                "Top": 855.5,
                "Height": 12,
                "Width": 38
              },
              {
                "WordText": "of",
                "Left": 312,
                "Top": 855.5,
                "Height": 9.5,
                "Width": 11
              },
              {
                "WordText": "50+",
                "Left": 327,
                "Top": 856,
                "Height": 9,
                "Width": 19.5
              },
              {
                "WordText": "cities",
                "Left": 351,
                "Top": 856,
                "Height": 9,
                "Width": 27.5
              },
              {
                "WordText": "in",
                "Left": 383.5,
                "Top": 856,
                "Height": 9,
                "Width": 9
              },
              {
                "WordText": "Spain",
                "Left": 397.5,
                "Top": 856,
                "Height": 11.5,
                "Width": 28.5
              },
              {
                "WordText": "(where",
                "Left": 431.5,
                "Top": 855,
                "Height": 12.5,
                "Width": 37
              },
              {
                "WordText": "geo-targeted",
                "Left": 473,
                "Top": 855.5,
                "Height": 12,
                "Width": 69.5
              },
              {
                "WordText": "ads",
                "Left": 548,
                "Top": 855.5,
                "Height": 9.5,
                "Width": 18
              },
              {
                "WordText": "were",
                "Left": 570.5,
                "Top": 858,
                "Height": 7,
                "Width": 26.5
              },
              {
                "WordText": "possible)",
                "Left": 602,
                "Top": 855,
                "Height": 12.5,
                "Width": 47
              },
              {
                "WordText": "based",
                "Left": 654.5,
                "Top": 855.5,
                "Height": 9.5,
                "Width": 30.5
              },
              {
                "WordText": "on",
                "Left": 690,
                "Top": 858,
                "Height": 7,
                "Width": 13
              },
              {
                "WordText": "their",
                "Left": 708,
                "Top": 855.5,
                "Height": 9.5,
                "Width": 25.5
              },
              {
                "WordText": "sales-",
                "Left": 738,
                "Top": 855.5,
                "Height": 9.5,
                "Width": 30
              }
            ],
            "MaxHeight": 12.5,
            "MinTop": 855
          },
          {
            "LineText": "wise similarity over time, using wavelets (and R).",
            "Words": [
              {
                "WordText": "wise",
                "Left": 228,
                "Top": 872,
                "Height": 9.5,
                "Width": 24
              },
              {
                "WordText": "similarity",
                "Left": 255.5,
                "Top": 872,
                "Height": 12,
                "Width": 49.5
              },
              {
                "WordText": "over",
                "Left": 308.5,
                "Top": 874.5,
                "Height": 7,
                "Width": 24
              },
              {
                "WordText": "time,",
                "Left": 335.5,
                "Top": 872,
                "Height": 11.5,
                "Width": 27.5
              },
              {
                "WordText": "using",
                "Left": 367.5,
                "Top": 872,
                "Height": 12,
                "Width": 27.5
              },
              {
                "WordText": "wavelets",
                "Left": 398.5,
                "Top": 872,
                "Height": 9.5,
                "Width": 47.5
              },
              {
                "WordText": "(and",
                "Left": 450,
                "Top": 871.5,
                "Height": 12.5,
                "Width": 23
              },
              {
                "WordText": "R).",
                "Left": 477,
                "Top": 871.5,
                "Height": 12.5,
                "Width": 13.5
              }
            ],
            "MaxHeight": 12.5,
            "MinTop": 871.5
          },
          {
            "LineText": "MANAGEMENT - SALES (Electrical Eng.)",
            "Words": [
              {
                "WordText": "MANAGEMENT",
                "Left": 49,
                "Top": 895,
                "Height": 9,
                "Width": 90.5
              },
              {
                "WordText": "-",
                "Left": 143,
                "Top": 899.5,
                "Height": 1,
                "Width": 7.5
              },
              {
                "WordText": "SALES",
                "Left": 154,
                "Top": 895,
                "Height": 9,
                "Width": 34.5
              },
              {
                "WordText": "(Electrical",
                "Left": 193.5,
                "Top": 894,
                "Height": 12,
                "Width": 56.5
              },
              {
                "WordText": "Eng.)",
                "Left": 255,
                "Top": 894,
                "Height": 12,
                "Width": 28.5
              }
            ],
            "MaxHeight": 12,
            "MinTop": 894
          },
          {
            "LineText": "Feb. 2009 - Aug. 2013",
            "Words": [
              {
                "WordText": "Feb.",
                "Left": 49,
                "Top": 917,
                "Height": 10.5,
                "Width": 23.5
              },
              {
                "WordText": "2009",
                "Left": 77.5,
                "Top": 918,
                "Height": 9.5,
                "Width": 28.5
              },
              {
                "WordText": "-",
                "Left": 110,
                "Top": 922.5,
                "Height": 1.5,
                "Width": 7.5
              },
              {
                "WordText": "Aug.",
                "Left": 121,
                "Top": 918,
                "Height": 12,
                "Width": 25.5
              },
              {
                "WordText": "2013",
                "Left": 151.5,
                "Top": 918,
                "Height": 9.5,
                "Width": 28
              }
            ],
            "MaxHeight": 12,
            "MinTop": 917
          },
          {
            "LineText": "Head of Sales, Spain & Portugal — Test &Measurement dept.",
            "Words": [
              {
                "WordText": "Head",
                "Left": 205,
                "Top": 917,
                "Height": 10,
                "Width": 30
              },
              {
                "WordText": "of",
                "Left": 239.5,
                "Top": 917,
                "Height": 10,
                "Width": 12
              },
              {
                "WordText": "Sales,",
                "Left": 255.5,
                "Top": 917,
                "Height": 12.5,
                "Width": 33.5
              },
              {
                "WordText": "Spain",
                "Left": 293.5,
                "Top": 917,
                "Height": 12.5,
                "Width": 32.5
              },
              {
                "WordText": "&",
                "Left": 330.5,
                "Top": 917.5,
                "Height": 10,
                "Width": 9.5
              },
              {
                "WordText": "Portugal",
                "Left": 344.5,
                "Top": 917,
                "Height": 12.5,
                "Width": 50
              },
              {
                "WordText": "—",
                "Left": 398.5,
                "Top": 922.5,
                "Height": 1.5,
                "Width": 7.5
              },
              {
                "WordText": "Test",
                "Left": 409.5,
                "Top": 918,
                "Height": 9.5,
                "Width": 24.5
              },
              {
                "WordText": "&Measurement",
                "Left": 438.5,
                "Top": 917,
                "Height": 10.5,
                "Width": 93
              },
              {
                "WordText": "dept.",
                "Left": 536,
                "Top": 917,
                "Height": 13,
                "Width": 30
              }
            ],
            "MaxHeight": 13,
            "MinTop": 917
          },
          {
            "LineText": "YOKOGAWA",
            "Words": [
              {
                "WordText": "YOKOGAWA",
                "Left": 204,
                "Top": 935.5,
                "Height": 9.5,
                "Width": 73.5
              }
            ],
            "MaxHeight": 9.5,
            "MinTop": 935.5
          },
          {
            "LineText": "• Applied analysis of sales and market trends to decide the direction of the department.",
            "Words": [
              {
                "WordText": "•",
                "Left": 217,
                "Top": 955.5,
                "Height": 4.5,
                "Width": 4.5
              },
              {
                "WordText": "Applied",
                "Left": 228,
                "Top": 952.5,
                "Height": 12,
                "Width": 40.5
              },
              {
                "WordText": "analysis",
                "Left": 273,
                "Top": 952.5,
                "Height": 12,
                "Width": 41.5
              },
              {
                "WordText": "of",
                "Left": 318.5,
                "Top": 952.5,
                "Height": 9.5,
                "Width": 11
              },
              {
                "WordText": "sales",
                "Left": 332.5,
                "Top": 952.5,
                "Height": 9.5,
                "Width": 26
              },
              {
                "WordText": "and",
                "Left": 362,
                "Top": 952.5,
                "Height": 9.5,
                "Width": 19
              },
              {
                "WordText": "market",
                "Left": 386,
                "Top": 952.5,
                "Height": 9.5,
                "Width": 37.5
              },
              {
                "WordText": "trends",
                "Left": 427,
                "Top": 952.5,
                "Height": 9.5,
                "Width": 34.5
              },
              {
                "WordText": "to",
                "Left": 465,
                "Top": 953.5,
                "Height": 8.5,
                "Width": 11
              },
              {
                "WordText": "decide",
                "Left": 479.5,
                "Top": 952.5,
                "Height": 9.5,
                "Width": 35.5
              },
              {
                "WordText": "the",
                "Left": 518.5,
                "Top": 952.5,
                "Height": 9.5,
                "Width": 17.5
              },
              {
                "WordText": "direction",
                "Left": 540,
                "Top": 952.5,
                "Height": 9.5,
                "Width": 47
              },
              {
                "WordText": "of",
                "Left": 591.5,
                "Top": 952.5,
                "Height": 9.5,
                "Width": 11
              },
              {
                "WordText": "the",
                "Left": 605,
                "Top": 952.5,
                "Height": 9.5,
                "Width": 18
              },
              {
                "WordText": "department.",
                "Left": 626.5,
                "Top": 952.5,
                "Height": 12,
                "Width": 67
              }
            ],
            "MaxHeight": 12,
            "MinTop": 952.5
          },
          {
            "LineText": "• Led a team of 7 people.",
            "Words": [
              {
                "WordText": "•",
                "Left": 217,
                "Top": 972,
                "Height": 4,
                "Width": 4.5
              },
              {
                "WordText": "Led",
                "Left": 229,
                "Top": 969,
                "Height": 9,
                "Width": 17.5
              },
              {
                "WordText": "a",
                "Left": 250.5,
                "Top": 971.5,
                "Height": 6.5,
                "Width": 5.5
              },
              {
                "WordText": "team",
                "Left": 260,
                "Top": 970,
                "Height": 8,
                "Width": 27
              },
              {
                "WordText": "of",
                "Left": 291,
                "Top": 969,
                "Height": 9,
                "Width": 11.5
              },
              {
                "WordText": "7",
                "Left": 305.5,
                "Top": 969.5,
                "Height": 8.5,
                "Width": 5.5
              },
              {
                "WordText": "people.",
                "Left": 315.5,
                "Top": 969,
                "Height": 11.5,
                "Width": 39
              }
            ],
            "MaxHeight": 11.5,
            "MinTop": 969
          },
          {
            "LineText": "Madrid, Spain",
            "Words": [
              {
                "WordText": "Madrid,",
                "Left": 693.5,
                "Top": 934.5,
                "Height": 11.5,
                "Width": 41.5
              },
              {
                "WordText": "Spain",
                "Left": 739,
                "Top": 935,
                "Height": 11.5,
                "Width": 28.5
              }
            ],
            "MaxHeight": 11.5,
            "MinTop": 934.5
          }
        ],
        "HasOverlay": 'true',
        "Message": "Total lines: 96"
      },
      "FileParseExitCode": 1,
      "TextOrientation": "0",
      "ParsedText": "Juan Jose Carin\tMountain View, CA 94041\t\r\n650-336-4590 | juanjose.carin@gmail.com\t\r\nData Scientist\tlinkedin.com/in/juanjosecarin I\tjuanjocarin.github.io\t\r\nProfessional Profile\t\r\nPassionate about data analysis and experiments, mainly focused on user behavior, experience, and engagement, with a solid\t\r\nbackground in data science and statistics, and extensive experience using data insights to drive business growth.\t\r\nEducation\t\r\n2016\tUniversity of California, Berkeley\tMaster of Information and Data Science\tGPA: 3.93\t\r\nRelevant courses:\t• Field Experiments\t•\tData Visualization and\t\r\n• Machine Learning\t• Applied Regression and Time Series\tCommunication\t\r\n• Machine Learning at Scale\tAnalysis\t•\tResearch Design and Applications for\t\r\n• Storing and Retrieving Data\t• Exploring and Analyzing Data\tData Analysis\t\r\n2014\tUniversidad Politécnica de Madrid\tM.S. in Statistical and Computational Information Processing\tGPA: 3.69\t\r\nRelevant courses:\t• Neural Networks and Statistical\t• Monte Carlo Techniques\t\r\n• Data Mining\tLearning\t• Numerical Methods in Finance\t\r\n• Multivariate Analysis\t• Regression and Prediction Methods\t• Stochastic Models in Finance\t\r\n• Time Series\t• Optimization Techniques\t• Bayesian Networks\t\r\n2005\tUniversidad Politécnica de Madrid\tM.S. in Telecommunication Engineering\tGPA: 3.03\t\r\nFocus Area:\tRadio communication systems (radar and mobile).\t\r\nFellowship:\tFirst year at University, due to Honors obtained last year at high school.\t\r\nSkills\t\r\nProgramming / Statistics\tBig Data\tVisualization\tOthers\t\r\nProficient:\tR, Python, SQL\tHadoop, Hive, MrJob\tTableau\tGit, A WS\t\r\nIntermediate:\tspss, SAS, Matiab\tSpark, Storm\tBash\t\r\nBasic:\tE Views, Demetra+\tD3.js\tGephi, Ne04j, QGIS\t\r\nExperience\t\r\nDATA SCIENCE\t\r\nJan. 2016- Mar. 2016\tData Scientist\t\r\nCONENTO\tMadrid, Spain (working remotely)\t\r\n• Designed and implemented the ETL pipeline for a predictive model of traffic on the main roads in\t\r\neastern Spain (a project for the Spanish government).\t\r\n• Automated scripts in R to extract, transform, clean (incl. anomaly detection), and load into MySQL\t\r\ndata from multiple data sources: road traffic sensors, accidents, road works, weather.\t\r\nJun. 2014 -Sep. 2014\tData Scientist\t\r\nCONENTO\tMadrid, Spain\t\r\n• Designed an experiment for Google Spain (conducted in October 2014) to measure the impact of\t\r\nYouTube ads on the sales of a car manufacturer's dealer network.\t\r\n• A matched-pair, cluster-randomized design, which involved selecting the test and control groups\t\r\nfrom a sample of 50+ cities in Spain (where geo-targeted ads were possible) based on their sales-\t\r\nwise similarity over time, using wavelets (and R).\t\r\nMANAGEMENT - SALES (Electrical Eng.)\t\r\nFeb. 2009 - Aug. 2013\tHead of Sales, Spain & Portugal — Test &Measurement dept.\t\r\nYOKOGAWA\tMadrid, Spain\t\r\n• Applied analysis of sales and market trends to decide the direction of the department.\t\r\n• Led a team of 7 people.\t\r\n",
      "ErrorMessage": "",
      "ErrorDetails": ""
    },
  ],
  "OCRExitCode": 1,
  "IsErroredOnProcessing": 'false',
  "ProcessingTimeInMilliseconds": 4.375,
  "SearchablePDFURL": "Searchable PDF not generated as it was not requested."
}
data1=[]
data = ocr_text['ParsedResults'][0]['Overlay']['Lines']

for i in range(0,len(data)):
    data1.append(data[i]['LineText'])

# for i in range(0,len(data)):
#     data1.append(data[i]['LineText']['LineText'])

# thelines= []
# for i in range(0,len(data1)):
#     thelines.append(data[i]['LineText'])


nlp = spacy.load("en_core_web_sm")
#sentences = ocr_text.split('•')
#both words, experience and Work experience
ind= data1.index('Experience')

mylist = data1[ind: ]
sentences = mylist

# for sent in sentences:
#     print(sent)
#     print('******')

has_subject = False
has_verb = False
has_adverb = False
has_conjunction = False

def makepdf():
  doc.save(r"C:\Users\umaRf\PycharmProjects\djangopractise\OCR\algo4.docx")
  from docx2pdf import convert
  convert(r"C:\Users\umaRf\PycharmProjects\djangopractise\OCR\algo4.docx", r"C:\Users\umaRf\PycharmProjects\djangopractise\OCR\algo4.pdf")
def is_grammatically_complete(sentence):
    # Load the sentence into spaCy's pipeline
    doc = nlp(sentence)

    # Check if the sentence contains at least one subject and one verb
    for token in doc:
        #print(token.text, token.pos_, token.dep_)
        if token.dep_ == "nsubj" and token.head.pos_ == "VERB" or token.pos_ == "ADV" or token.pos_ == "VERB" or token.pos_ == "CCONJ" or token.pos_ == "SCONJ":
            # if len(sentence.split(' ')) > 2:
              return True

    return False

fulltext=[]
doc = docx.Document(r"cv_template.docx")
old_text = "Job desc"

# Add the heading "Experience" to the document
#doc.add_heading("Experience", level=1)

def with_grammeralgo():

  for sent in sentences:
      if is_grammatically_complete(sent):
          #print(sent)
          fulltext.append(sent+'\n')

          # Add each item from the experience list to the document
          #doc.add_paragraph(sent + '\n\n')
          #doc.add_paragraph('    ')
          # Save the Word document

def without_grammeralgo():
  for sent in sentences:
    fulltext.append(sent)
    # Add each item from the experience list to the document
    if(len(sent.split(' '))>2):
      #print(sent)
      doc.add_paragraph(sent)
      #doc.add_paragraph('    ')

#without_grammeralgo()
with_grammeralgo()




def replacedata(date,company, summary):
  doc = docx.Document(r"cv_template.docx")
  import locationtagger

  for paragraph in doc.paragraphs:
    if 'Job desc' in paragraph.text:
      # Replace the text while maintaining the same style and formatting
      for run in paragraph.runs:
        if 'Job desc' in run.text:
          new_text = run.text.replace("Job desc", ' '.join(fulltext))
          run.text = new_text
          break
          # break out of the loop once the text is replaced in one run object

    if 'Location' in paragraph.text:
      # Replace the text while maintaining the same style and formatting
      for run in paragraph.runs:
        if 'Location' in run.text:
          place_entity = locationtagger.find_locations(text=' '.join(fulltext))
          new_text = run.text.replace('Location', str(place_entity.regions[0])+ ', ' +str(place_entity.countries[0]))
          run.text = new_text
          break

    if 'Date' in paragraph.text:
        # Replace the text while maintaining the same style and formatting
        for run in paragraph.runs:
          if 'Date' in run.text:
            new_text = run.text.replace("Date", ' '.join(date))
            run.text = new_text
            break
    if 'Company' in paragraph.text:
        # Replace the text while maintaining the same style and formatting
        for run in paragraph.runs:
          if 'Company' in run.text:
            new_text = run.text.replace("Company", ' '.join(company))
            run.text = new_text
            break
    rep = 'Summary'
    if rep in paragraph.text:
        # Replace the text while maintaining the same style and formatting
        for run in paragraph.runs:
          if rep in run.text:
            new_text = run.text.replace(rep, ' '.join(summary))
            run.text = new_text
            break


  doc.save("modified1.docx")
  convert(r'modified1.docx', r"modified.pdf")
  os.remove('modified1.docx')

replacedata(date='2016-2020', company= 'Amazon', summary='Hey I am softeare dev with sound knowledge of'
                                                         ' this and that and I am going to'
                                                         ' r and ohoigndg ')
# Save the document
#makepdf()
