
import nltk
ocr_text = ''' 

****** Result for Image/Page 1 ******
UMAR FAROOQ KHAN	
SOFTWARE DEVELOPER	
PERSONAL PROFILE	
A full-stack software developer with holistic knowledge of Android and Web	
CONTACT ME AT	
Development along with Python. Experienced in A1 and Data Science	
Techniques.	
umar.farooq407@gmail.com	
WORK EXPERIENCE	
https://www.linkedin.com/in/u	
Android Developer (Remote Location)	
mar-farooq-khan-107494167/	
Lesptitsmonstres I Montpellier, France I Aug 2019- February 2021	
+923365297082	• App Development in Native Android Platform with Android SDK/Java.	
• Using the Google Firebase (No-SQL database) and XML layout.	
• Manual Quality Assurance (QA), Maintainance, and Scaling.	
12-12-1997	
Android Developer (Contract, Remote Location)	
PROGRAMMING	Bothofus I Stockholm, Sweden I September 2020	
LANGUAGES	
Figma UI/UX Mockup Translation into Android screens by XML with	
PYTHON	
functionality.	
ANDROID	
JAVA	• Firebase Push Notification and Rest API implementation for server calls.	
K O TLIN	
JAVASCRIPT/ TYPESCRIPT	• Implementation of logo animation with Java/kotlin.	
HTML	
Kotlin Developer Intern	
E-conceptions I Islamabad, Pakistan I June 2019- July 2019	
• Support for basic programming: abstraction, inheritance, exception	
CLIENT-SIDE	
ANDROID DEVELOPMENT	handling, and polymorphism using Kotlin.	
ANGULAR/ NODE J s	
• Support for implementation of lottery-checker automation tool.	
FLASK	
Social Media Forensics, Data Science Intern	
DATABASE	
FIREBASE	National Center for Cyber Security I Islamabad I May 2019-July 2019	
MICROSOFT SQL SERVER	
MYSQL	• Development of Regional Database with Twitter API tweepy using Python.	
• Manual Labelling of the dataset into multi-label and multi-classes.	
• Implementation of exploration of trend (hashtag) origination on social	
ACHEIVEMENTS	media using GPS location information present in dataset.	
LEVEL 2 SELLER ON FIVERR	
LEVEL 1 SELLER ON FIVERR	Sentiment Analysis on Local Regional Data of Twitter with Machine	
Learning ( Bachelor's Thesis )	
• A project backed by the National Center for Cyber Security (NCCS), an	
agency commissioned by the Government of Pakistan.	
• GetOldTweets python API for fetching Twitter data.	
• Preprocessing of the dataset using Regex and Pandas Python Library.	
• Implementation of SVM and Naive Bayes algorithms in Python with the	
Sklearn and Machine Learning (ML) algorithms with an accuracy of 97%.	
• Front-End on Flask and the Back-End on Python.	

****** Result for Image/Page 2 ******
PROJECTS	
4	
Job Portal Website: Job portal website made from Django web	
framework	
•	Pulsar: A news aggregator website using Angular & Nodejs.	
Tailored/customized newsfeed using ML and A1 algorithms to show	
related news from different news sources.	
CONTACT ME AT	
•	Career Counseling Platform: A career counseling website made from	
umar.farooq407@gmail.com	
Angular & Nodejs where students, job seekers, and one who wants to	
switch their career can seek guidance from top leading career	
https://www.linkedin.com/in/u	
counselling coaches.	
mar-farooq-khan-107494167/	
• Blood Donation App: An android app for Urgent Blood Donation and for	
+923365297082	
Finding Donors.	
12-12-1997	
News journal priorities on UN SGD: Extensive work on web Scraping and	
Detailed Sentiment Analysis and comparison of priorities of UN SGD	
goals on 3 top American News journals	
Spanish Tweets Sentiment Analysis: Spanish tweets were scraped by	
the T weepy API of some political parties and then classified and	
analyzed the sentiments.	
What's My Profit: E-commerce Income Calculator: Automation tool that	
calculates various income sources files and converts sales into local	
currency at an exchange rate	
EDUCATIONAL HISTORY	
Bachelor of Science in Computer Science (BSCS)	
Air University, Islamabad | 2016-2020	
• Relevant Subjects: Data Science, Information & Network Security,	
Android Development, Full Stack Web Development (MEAN), Digital	
Image Processing, Human-Machine Interaction (HMI), Software	
Engineering, Visual Programming C#, Data Structure & Algorithms, Intro	
to A1, Computer Architecture, Operating Systems (Linux Environment),	
Assembly Language	
Pre-Engineering, High School	
ICB G-6/3 1 Islamabad, Pakistan 12013-2015	
LANGUAGES	
• English (Fluent)	
• Urdu(Fluent)	
EXTRACURRICULAR ACTIVITIES	
Cricket	
Writing Programming Articles (Medium.com)	
YouTube Content Creation, Cinematography_	
Photography, Travelling	

'''
import spacy
import docx
import re
nlp = spacy.load("en_core_web_sm")
pattern = r"\•\s.+"

# Filter the sentences that match the pattern
bullet_sentences = re.findall(pattern, ocr_text)

# Print the filtered sentences
for s in bullet_sentences:
    print(s)

sentences = bullet_sentences
# for sent in sentences:
#     print(sent)
#     print('******')

has_subject = False
has_verb = False
has_adverb = False
has_conjunction = False

def is_grammatically_complete(sentence):
    # Load the sentence into spaCy's pipeline
    doc = nlp(sentence)

    # Check if the sentence contains at least one subject and one verb
    for token in doc:
        #print(token.text, token.pos_, token.dep_)
        if token.dep_ == "nsubj" and token.head.pos_ == "VERB" or token.pos_ == "ADV" or token.pos_ == "VERB" or token.pos_ == "CCONJ" or token.pos_ == "SCONJ":
            return True

    return False

fulltext=[]
doc = docx.Document()

# Add the heading "Experience" to the document
doc.add_heading("Experience", level=1)
for sent in sentences:
    if is_grammatically_complete(sent):
        print('yes', sent)
        print('********')
        fulltext.append(sent)
        # Add each item from the experience list to the document
        doc.add_paragraph(sent + '\n')
        doc.add_paragraph('    ')
        # Save the Word document

doc.save("experience3.docx")
from docx2pdf import convert
convert("experience3.docx", "output1.pdf")