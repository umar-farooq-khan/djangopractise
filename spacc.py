import PyPDF2
import spacy
import docx
from docx2pdf import convert

# Load the PDF document and extract text
pdf_file = open(r"C:\Users\umaRf\OneDrive\Desktop\samplecvs\resume_juanjosecarin.pdf", 'rb')

pdf_reader = PyPDF2.PdfReader(pdf_file)
text = ""
for page in pdf_reader.pages:
    text += page.extract_text()

# Load the spaCy English language model
nlp = spacy.load("en_core_web_sm")
print(text)
# Process the text with spaCy
doc = nlp(text)

# Extract all experience entities
experience = []
for ent in doc.ents:
    if ent.label_ == "DATE":
        if ent.start != 0 and doc[ent.start - 1].text.lower() in ["from", "since"]:
            experience.append(ent.text)
        elif ent.end != len(doc) and doc[ent.end].text.lower() == "present":
            experience.append(ent.text)


# Print the extracted experience entities

# Create a new Word document
doc = docx.Document()

# Add the heading "Experience" to the document
doc.add_heading("Experience", level=1)

# Add each item from the experience list to the document
x= text.split('Experience')
doc.add_paragraph(x[1])

# Save the Word document
doc.save("experience.docx")

# Convert a single file from docx to pdf
convert("experience.docx", "output.pdf")