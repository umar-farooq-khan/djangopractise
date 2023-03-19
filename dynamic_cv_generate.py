import docx
import os

def replace_text_in_document(lists, template_doc_path, output_doc_path):
    # Create a dictionary to hold the values to be replaced
    text_dict = {}
    num_jobs = len(lists['companylist'])

    for i in range(num_jobs):
        key1 = f'Company{i+1}'
        value1 = str(lists['companylist'][0][i])
        text_dict[key1] = value1

        key2 = f'desc{i+1}'
        value2 = lists['job_desc_list'][0][i]
        text_dict[key2] = value2

    # Load the template document
    doc = docx.Document(template_doc_path)

    # Replace the text in the document
    for old_text, new_text in text_dict.items():
        for paragraph in doc.paragraphs:
            if old_text in paragraph.text:
                for run in paragraph.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
                        break

    # Save the modified document
    doc.save(output_doc_path)

    # Convert the document to PDF and remove the temporary DOCX file
    convert(output_doc_path, output_doc_path[:-4] + '.pdf')
    os.remove(output_doc_path)

def convert(docx_path, pdf_path):
    # Install PyPDF2 and import it
    !pip install PyPDF2
    import PyPDF2

    # Open the DOCX file and create a PDF writer object
    docx_file = open(docx_path, 'rb')
    pdf_writer = PyPDF2.PdfFileWriter()

    # Open the DOCX file with python-docx and convert each page to a PDF page
    doc = docx.Document(docx_file)
    for page in doc.sections:
        pdf_page = PyPDF2.pdf.PageObject.createBlankPage(None, page.width, page.height)
        pdf_writer.addPage(pdf_page)

    # Write the PDF to disk
    pdf_output_file = open(pdf_path, 'wb')
    pdf_writer.write(pdf_output_file)

    # Clean up
    pdf_output_file.close()
    docx_file.close()

# Example usage
lists = {
    'companylist': [['ABC', 'DEF', 'GHI', 'JKL', 'MNO']],
    'job_fromdatelist': [['01/2015', '04/2017', '07/2018', '08/2020', '02/2022']],
    'job_todatelist': [['03/2017', '06/2018', '08/2020', '02/2022', '']],
    'job_desc_list': [['Job 1 description', 'Job 2 description', 'Job 3 description', 'Job 4 description', 'Job 5 description']]
}

replace_text_in_document(lists, 'cv_new.docx', 'output.docx')
